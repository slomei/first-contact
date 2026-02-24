"""Tests for DOCX and XLSX document creation."""

import os
from unittest.mock import patch

import pytest

import documents
import memory


class TestCreateDocx:
    def test_creates_file(self, isolated_env):
        path = documents.create_docx("Hello world", "test.docx")
        assert os.path.exists(path)
        assert path.endswith(".docx")

    def test_valid_docx(self, isolated_env):
        """File can be opened by python-docx."""
        from docx import Document
        path = documents.create_docx("Hello world", "test.docx")
        doc = Document(path)
        assert len(doc.paragraphs) >= 1

    def test_paragraph_splitting(self, isolated_env):
        from docx import Document
        content = "First paragraph.\n\nSecond paragraph.\n\nThird paragraph."
        path = documents.create_docx(content, "test.docx")
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs if p.text]
        assert len(texts) == 3

    def test_title_as_heading(self, isolated_env):
        from docx import Document
        path = documents.create_docx("Body text", "test.docx", title="My Title")
        doc = Document(path)
        # First paragraph should be the heading
        assert doc.paragraphs[0].text == "My Title"
        assert doc.paragraphs[0].style.name.startswith("Heading")

    def test_no_title_no_heading(self, isolated_env):
        from docx import Document
        path = documents.create_docx("Body text", "test.docx")
        doc = Document(path)
        # First paragraph should be body text, not a heading
        assert doc.paragraphs[0].text == "Body text"
        assert not doc.paragraphs[0].style.name.startswith("Heading")

    def test_fallback_to_txt(self, isolated_env):
        with patch.object(documents, "_DOCX_AVAILABLE", False):
            path = documents.create_docx("Hello world", "test.docx", title="Title")
            assert path.endswith(".txt")
            assert os.path.exists(path)
            content = open(path).read()
            assert "Title" in content
            assert "Hello world" in content


class TestCreateXlsx:
    def test_creates_file(self, isolated_env):
        data = [["a", "b"], ["c", "d"]]
        path = documents.create_xlsx(data, "test.xlsx")
        assert os.path.exists(path)
        assert path.endswith(".xlsx")

    def test_valid_xlsx(self, isolated_env):
        """File can be opened by openpyxl."""
        from openpyxl import load_workbook
        data = [["a", "b"], ["c", "d"]]
        path = documents.create_xlsx(data, "test.xlsx")
        wb = load_workbook(path)
        ws = wb.active
        assert ws.cell(1, 1).value == "a"
        assert ws.cell(2, 2).value == "d"

    def test_data_rows(self, isolated_env):
        from openpyxl import load_workbook
        data = [[1, 2], [3, 4], [5, 6]]
        path = documents.create_xlsx(data, "test.xlsx")
        wb = load_workbook(path)
        ws = wb.active
        assert ws.max_row == 3

    def test_headers_bold(self, isolated_env):
        from openpyxl import load_workbook
        data = [["x", "y"]]
        path = documents.create_xlsx(data, "test.xlsx", headers=["Col A", "Col B"])
        wb = load_workbook(path)
        ws = wb.active
        assert ws.cell(1, 1).value == "Col A"
        assert ws.cell(1, 1).font.bold
        # Data starts at row 2
        assert ws.cell(2, 1).value == "x"

    def test_no_headers_data_at_row1(self, isolated_env):
        from openpyxl import load_workbook
        data = [["x", "y"]]
        path = documents.create_xlsx(data, "test.xlsx")
        wb = load_workbook(path)
        ws = wb.active
        assert ws.cell(1, 1).value == "x"

    def test_column_widths_adjusted(self, isolated_env):
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter
        data = [["short", "a much longer cell value here"]]
        path = documents.create_xlsx(data, "test.xlsx")
        wb = load_workbook(path)
        ws = wb.active
        width_a = ws.column_dimensions[get_column_letter(1)].width
        width_b = ws.column_dimensions[get_column_letter(2)].width
        assert width_b > width_a

    def test_fallback_to_csv(self, isolated_env):
        with patch.object(documents, "_XLSX_AVAILABLE", False):
            data = [["a", "b"], ["c", "d"]]
            path = documents.create_xlsx(data, "test.xlsx", headers=["H1", "H2"])
            assert path.endswith(".csv")
            assert os.path.exists(path)
            content = open(path).read()
            assert "H1" in content
            assert "a" in content


class TestToolDispatch:
    def test_create_docx_dispatch(self, isolated_env):
        import tools
        result, is_error = tools.execute_tool(
            "create_docx",
            {"content": "Test content", "filename": "out.docx"},
        )
        assert not is_error
        assert "out.docx" in result
        # File should exist at the returned path
        path = result.split(": ", 1)[1]
        assert os.path.exists(path)

    def test_create_xlsx_dispatch(self, isolated_env):
        import tools
        result, is_error = tools.execute_tool(
            "create_xlsx",
            {"data": [["a", "b"]], "filename": "out.xlsx"},
        )
        assert not is_error
        assert "out.xlsx" in result
        path = result.split(": ", 1)[1]
        assert os.path.exists(path)
